"""Тесты фичи 1 — загрузка входных файлов-вложений.

Покрывает: модель/репозиторий attachments, извлечение текста (успех/деградация,
кириллица), подачу текста в слой A, правило удаления-до-использования,
пометку used_in_context при сборке контекста и REST-эндпоинты.

Кириллица проверяется на .txt/.docx (шрифто-независимо). PDF-путь проверяется
на ASCII через reportlab (стандартный шрифт Helvetica) — это не зависит от
системных шрифтов и надёжно работает на всей CI-матрице.
"""

from __future__ import annotations

import io
from pathlib import Path

from fastapi.testclient import TestClient

from pov_generator.application.attachment_service import AttachmentService, _sanitize_filename
from pov_generator.application.context_service import ContextService
from pov_generator.application.planning_service import PlanningService
from pov_generator.application.project_service import ProjectService
from pov_generator.application.registry_service import RegistryService
from pov_generator.common.errors import ConflictError
from pov_generator.domain.registry import ObjectRef
from pov_generator.infrastructure.filesystem_registry import FilesystemRegistryLoader
from pov_generator.infrastructure.sqlite_runtime import SqliteRuntime
from pov_generator.interfaces.api import create_app

REPO_ROOT = Path(__file__).resolve().parents[1]
OBJECTIVE_REF = "common.requirements_specification@1.0.0"


def _bootstrap(tmp_path: Path, *, expand: bool = False):
    registry_service = RegistryService(FilesystemRegistryLoader(REPO_ROOT / "templates"))
    runtime = SqliteRuntime()
    project_service = ProjectService(runtime)
    snapshot, report = registry_service.validate()
    assert report.is_valid
    workspace = tmp_path / "case"
    bootstrap = project_service.init_project(
        workspace=workspace,
        name="attach test",
        objective_ref=ObjectRef.parse(OBJECTIVE_REF),
        request_text="Нужна CRM-интеграция для отдела продаж.",
        domain_packs=(),
    )
    if expand:
        PlanningService(runtime).expand_graph(workspace, snapshot)
    return workspace, bootstrap.manifest.project_id, runtime, registry_service


def _make_pdf(text: str) -> bytes:
    """Минимальный одностраничный PDF с ASCII-текстом (Helvetica)."""
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas

    buffer = io.BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=A4)
    pdf.setFont("Helvetica", 14)
    pdf.drawString(72, 760, text)
    pdf.save()
    return buffer.getvalue()


def _make_docx(*paragraphs: str) -> bytes:
    from docx import Document

    document = Document()
    for paragraph in paragraphs:
        document.add_paragraph(paragraph)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


# --- извлечение текста + слой A ----------------------------------------------


def test_text_upload_extracts_and_feeds_layer_a(tmp_path: Path) -> None:
    workspace, project_id, runtime, _ = _bootstrap(tmp_path)
    service = AttachmentService(runtime)

    record = service.upload(
        workspace,
        project_id,
        filename="требования.txt",
        content="Требования клиента: интеграция CRM.".encode("utf-8"),
        mime_type="text/plain",
        extract_in_background=False,
    )
    stored = runtime.load_attachment(workspace, record.attachment_id)
    assert stored.extraction_status == "succeeded"
    assert stored.linked_position_id == f"attachment.{record.attachment_id}"

    knowledge = runtime.load_knowledge(workspace)
    position = knowledge.positions[stored.linked_position_id]
    assert position.type == "fact"
    assert position.source == "input"
    assert "Требования клиента" in position.statement


def test_docx_cyrillic_extraction(tmp_path: Path) -> None:
    workspace, project_id, runtime, _ = _bootstrap(tmp_path)
    service = AttachmentService(runtime)
    record = service.upload(
        workspace,
        project_id,
        filename="brief.docx",
        content=_make_docx("Заголовок", "Параграф с кириллицей: цели проекта."),
        extract_in_background=False,
    )
    stored = runtime.load_attachment(workspace, record.attachment_id)
    assert stored.extraction_status == "succeeded"
    text = (workspace / stored.extracted_text_ref).read_text(encoding="utf-8")
    assert "кириллицей" in text


def test_docx_table_extraction(tmp_path: Path) -> None:
    """Регрессия: таблицы .docx раньше терялись (читались только параграфы).

    Документ с параграфом + таблицей: в извлечённом тексте должны быть и
    параграф, и содержимое всех ячеек таблицы (в т.ч. кириллица)."""
    from docx import Document

    document = Document()
    document.add_paragraph("Вводный параграф документа.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Параметр"
    table.cell(0, 1).text = "Значение"
    table.cell(1, 0).text = "Срок поставки"
    table.cell(1, 1).text = "30 рабочих дней"
    buffer = io.BytesIO()
    document.save(buffer)

    workspace, project_id, runtime, _ = _bootstrap(tmp_path)
    service = AttachmentService(runtime)
    record = service.upload(
        workspace,
        project_id,
        filename="contract.docx",
        content=buffer.getvalue(),
        extract_in_background=False,
    )
    stored = runtime.load_attachment(workspace, record.attachment_id)
    assert stored.extraction_status == "succeeded"
    text = (workspace / stored.extracted_text_ref).read_text(encoding="utf-8")
    assert "Вводный параграф документа." in text
    for cell_value in ("Параметр", "Значение", "Срок поставки", "30 рабочих дней"):
        assert cell_value in text, f"ячейка таблицы потеряна: {cell_value!r}"


def test_pdf_extraction(tmp_path: Path) -> None:
    workspace, project_id, runtime, _ = _bootstrap(tmp_path)
    service = AttachmentService(runtime)
    record = service.upload(
        workspace,
        project_id,
        filename="spec.pdf",
        content=_make_pdf("Project requirements and scope."),
        mime_type="application/pdf",
        extract_in_background=False,
    )
    stored = runtime.load_attachment(workspace, record.attachment_id)
    assert stored.extraction_status == "succeeded"
    text = (workspace / stored.extracted_text_ref).read_text(encoding="utf-8")
    assert "requirements" in text


def test_unsupported_format_degrades(tmp_path: Path) -> None:
    workspace, project_id, runtime, _ = _bootstrap(tmp_path)
    service = AttachmentService(runtime)
    record = service.upload(
        workspace,
        project_id,
        filename="picture.png",
        content=b"\x89PNG\r\n\x1a\n",
        extract_in_background=False,
    )
    stored = runtime.load_attachment(workspace, record.attachment_id)
    assert stored.extraction_status == "unsupported"
    assert stored.linked_position_id is None
    # Файл всё равно скачивается.
    assert runtime.load_attachment_content(workspace, record.attachment_id)


def test_empty_pdf_degrades_to_failed(tmp_path: Path) -> None:
    workspace, project_id, runtime, _ = _bootstrap(tmp_path)
    service = AttachmentService(runtime)
    # Валидный PDF без текстового слоя (пустая страница) → failed, не падение.
    record = service.upload(
        workspace,
        project_id,
        filename="blank.pdf",
        content=_make_pdf(""),
        mime_type="application/pdf",
        extract_in_background=False,
    )
    stored = runtime.load_attachment(workspace, record.attachment_id)
    assert stored.extraction_status == "failed"
    assert stored.extraction_error


def test_empty_file_rejected(tmp_path: Path) -> None:
    workspace, project_id, runtime, _ = _bootstrap(tmp_path)
    service = AttachmentService(runtime)
    try:
        service.upload(workspace, project_id, filename="empty.txt", content=b"", extract_in_background=False)
        raise AssertionError("пустой файл должен отклоняться")
    except ConflictError:
        pass
    assert runtime.list_attachments(workspace) == []


def test_duplicate_upload_is_idempotent(tmp_path: Path) -> None:
    workspace, project_id, runtime, _ = _bootstrap(tmp_path)
    service = AttachmentService(runtime)
    content = "одинаковый текст файла".encode("utf-8")
    first = service.upload(workspace, project_id, filename="dup.txt", content=content, extract_in_background=False)
    second = service.upload(workspace, project_id, filename="dup.txt", content=content, extract_in_background=False)
    # Точный повтор (то же содержимое и имя) → та же запись, без дубля.
    assert second.attachment_id == first.attachment_id
    assert len(runtime.list_attachments(workspace)) == 1
    # То же содержимое под другим именем — намеренно отдельный файл.
    third = service.upload(workspace, project_id, filename="other.txt", content=content, extract_in_background=False)
    assert third.attachment_id != first.attachment_id
    assert len(runtime.list_attachments(workspace)) == 2


def test_binary_under_text_extension_degrades(tmp_path: Path) -> None:
    workspace, project_id, runtime, _ = _bootstrap(tmp_path)
    service = AttachmentService(runtime)
    # Бинарь (с NUL-байтами) под расширением .txt → failed, не «текст из мусора».
    record = service.upload(
        workspace,
        project_id,
        filename="fake.txt",
        content=b"PK\x03\x04\x00\x00binary\x00content",
        extract_in_background=False,
    )
    stored = runtime.load_attachment(workspace, record.attachment_id)
    assert stored.extraction_status == "failed"
    assert stored.extraction_error
    assert stored.linked_position_id is None
    # Файл всё равно скачивается.
    assert runtime.load_attachment_content(workspace, record.attachment_id)


# --- лимиты и гигиена --------------------------------------------------------


def test_sanitize_filename_strips_path_traversal() -> None:
    assert _sanitize_filename("../../etc/passwd") == "passwd"
    assert _sanitize_filename("a/b\\c.txt") == "c.txt"
    assert _sanitize_filename("кириллица .pdf") == "кириллица .pdf"
    assert _sanitize_filename("") == "file"


def test_size_limit_enforced(tmp_path: Path) -> None:
    workspace, project_id, runtime, _ = _bootstrap(tmp_path)
    service = AttachmentService(runtime, max_file_bytes=10)
    try:
        service.upload(workspace, project_id, filename="big.txt", content=b"x" * 11, extract_in_background=False)
        raise AssertionError("ожидалась ошибка лимита размера")
    except ConflictError:
        pass


def test_count_limit_enforced(tmp_path: Path) -> None:
    workspace, project_id, runtime, _ = _bootstrap(tmp_path)
    service = AttachmentService(runtime, max_files_per_project=1)
    service.upload(workspace, project_id, filename="a.txt", content=b"a", extract_in_background=False)
    try:
        service.upload(workspace, project_id, filename="b.txt", content=b"b", extract_in_background=False)
        raise AssertionError("ожидалась ошибка лимита количества")
    except ConflictError:
        pass


# --- удаление-до-использования -----------------------------------------------


def test_delete_before_use_removes_position(tmp_path: Path) -> None:
    workspace, project_id, runtime, _ = _bootstrap(tmp_path)
    service = AttachmentService(runtime)
    record = service.upload(
        workspace, project_id, filename="a.txt", content="данные".encode("utf-8"), extract_in_background=False
    )
    position_id = runtime.load_attachment(workspace, record.attachment_id).linked_position_id
    service.delete(workspace, record.attachment_id)

    stored = runtime.load_attachment(workspace, record.attachment_id)
    assert stored.is_deleted
    active_ids = [p.identifier for p in runtime.load_knowledge(workspace).active()]
    assert position_id not in active_ids
    # Файлы удалены с диска.
    assert not (workspace / stored.storage_path).exists()


def test_extract_does_not_resurrect_deleted_attachment(tmp_path: Path) -> None:
    """TOCTOU: если вложение удалили во время (долгого) извлечения, extract не
    должен «воскресить» его — ни Position в слой A, ни файл с текстом."""
    import uuid
    from hashlib import sha256

    from pov_generator.common.serialization import utc_now_iso
    from pov_generator.domain.attachments import AttachmentRecord

    workspace, project_id, runtime, _ = _bootstrap(tmp_path)
    service = AttachmentService(runtime)
    # Сохраняем pending-вложение напрямую (без извлечения), чтобы Position
    # ещё НЕ существовала к моменту гонки.
    content = "данные".encode("utf-8")
    attachment_id = str(uuid.uuid4())
    record = AttachmentRecord(
        attachment_id=attachment_id,
        project_id=project_id,
        original_filename="slow.txt",
        mime_type="text/plain",
        size_bytes=len(content),
        sha256=sha256(content).hexdigest(),
        storage_path=f"attachments/{attachment_id}.txt",
        extraction_status="pending",
        created_at=utc_now_iso(),
    )
    runtime.store_attachment(workspace, attachment=record, content=content)

    # Гонка: первый load (в начале extract) видит активное вложение, повторный
    # load (перед записью результата) — уже удалённое.
    original_load = runtime.load_attachment
    calls = {"n": 0}

    def racing_load(ws: Path, aid: str):
        calls["n"] += 1
        loaded = original_load(ws, aid)
        if calls["n"] >= 2:
            from dataclasses import replace as _replace

            return _replace(loaded, is_deleted=True)
        return loaded

    runtime.load_attachment = racing_load  # type: ignore[method-assign]
    try:
        service.extract(workspace, attachment_id)
    finally:
        runtime.load_attachment = original_load  # type: ignore[method-assign]

    # Ключевой инвариант: удалённое вложение не вернуло Position в слой A.
    knowledge = runtime.load_knowledge(workspace)
    assert f"attachment.{attachment_id}" not in [p.identifier for p in knowledge.active()]
    # extract вышел до записи результата — статус не перебит на succeeded.
    assert runtime.load_attachment(workspace, attachment_id).extraction_status == "pending"


def test_delete_after_use_is_forbidden(tmp_path: Path) -> None:
    workspace, project_id, runtime, _ = _bootstrap(tmp_path)
    service = AttachmentService(runtime)
    record = service.upload(workspace, project_id, filename="a.txt", content=b"data", extract_in_background=False)
    runtime.mark_attachment_used(workspace, record.attachment_id)
    try:
        service.delete(workspace, record.attachment_id)
        raise AssertionError("удаление использованного вложения должно быть запрещено")
    except ConflictError:
        pass


# --- used_in_context при сборке контекста ------------------------------------


def test_context_build_marks_attachment_used(tmp_path: Path) -> None:
    workspace, project_id, runtime, registry_service = _bootstrap(tmp_path, expand=True)
    service = AttachmentService(runtime)
    record = service.upload(
        workspace,
        project_id,
        filename="brief.txt",
        content="Дополнительные требования к проекту.".encode("utf-8"),
        extract_in_background=False,
    )
    assert runtime.load_attachment(workspace, record.attachment_id).used_in_context is False

    snapshot, _ = registry_service.validate()
    context_service = ContextService(runtime)
    # Берём leaf, который (а) потребляет вложения (attachments в raw_inputs —
    # иначе used_in_context не выставится) И (б) не требует upstream-артефактов
    # (иначе build_for_task упадёт на отсутствующем входе). Оба условия нужны:
    # «без обязательных входов» само по себе не гарантирует потребление вложения.
    leaf = next(
        t
        for t in runtime.list_tasks(workspace)
        if t.template_type == "leaf"
        and "attachments" in snapshot.resolve_template(t.template_ref).inputs.raw_inputs
        and not snapshot.resolve_template(t.template_ref).inputs.required_artifact_roles
    )
    context_service.build_for_task(workspace, snapshot, leaf.task_id)

    assert runtime.load_attachment(workspace, record.attachment_id).used_in_context is True


def test_oversized_source_trimmed_visibly_not_silently(tmp_path: Path, monkeypatch) -> None:
    """Не влезающий первоисточник усекается ВИДИМО (с пометкой + запись в аудит),
    а не режется молча и не роняет задачу."""
    workspace, project_id, runtime, registry_service = _bootstrap(tmp_path, expand=True)
    service = AttachmentService(runtime)
    service.upload(
        workspace,
        project_id,
        filename="big-brief.txt",
        content=("Очень длинный текст входного файла. " * 200).encode("utf-8"),
        extract_in_background=False,
    )
    snapshot, _ = registry_service.validate()
    # Прижимаем бюджет шаблона так, что источник не влезает целиком.
    monkeypatch.setenv("POV_TEMPLATE_CONTEXT_MAX_TOKENS", "2000")
    # Leaf, который потребляет вложения И не требует upstream-артефактов —
    # иначе build_for_task упал бы на отсутствующем обязательном входе
    # (например, scope_boundary_definition требует normalized_request).
    leaf = next(
        t
        for t in runtime.list_tasks(workspace)
        if t.template_type == "leaf"
        and "attachments" in snapshot.resolve_template(t.template_ref).inputs.raw_inputs
        and not snapshot.resolve_template(t.template_ref).inputs.required_artifact_roles
    )
    result = ContextService(runtime).build_for_task(
        workspace, snapshot, leaf.task_id, model_context_window=200_000
    )
    sources = [it for it in result.manifest.items if it.title == "Входной файл заказчика"]
    assert sources, "источник должен остаться (усечённым), а не исчезнуть"
    assert "усечён" in sources[0].content                                   # видимая пометка
    assert any("усечён" in note for note in result.manifest.excluded_items)  # в аудите


# --- REST API ----------------------------------------------------------------


def _api_client(tmp_path: Path) -> TestClient:
    app = create_app(repo_root=REPO_ROOT, runtime_root=tmp_path / "runtime")
    return TestClient(app)


def _create_project(client: TestClient) -> str:
    response = client.post(
        "/api/projects",
        json={
            "name": "api attach",
            "objective_ref": OBJECTIVE_REF,
            "request_text": "CRM интеграция.",
            "domain_pack_refs": ["frontend.web_workspace@1.0.0"],
        },
    )
    assert response.status_code == 200, response.text
    return response.json()["project_id"]


def test_api_attachment_crud(tmp_path: Path) -> None:
    client = _api_client(tmp_path)
    project_id = _create_project(client)

    # upload
    upload = client.post(
        f"/api/projects/{project_id}/attachments",
        files={"file": ("brief.txt", b"Tekst zadania.", "text/plain")},
    )
    assert upload.status_code == 200, upload.text
    attachment_id = upload.json()["attachment_id"]

    # list (extraction is synchronous-enough for txt, but background thread may
    # race; the record exists regardless of extraction status).
    listing = client.get(f"/api/projects/{project_id}/attachments")
    assert listing.status_code == 200
    items = listing.json()
    assert any(item["attachment_id"] == attachment_id for item in items)

    # download
    download = client.get(f"/api/projects/{project_id}/attachments/{attachment_id}/download")
    assert download.status_code == 200
    assert download.content == b"Tekst zadania."
    assert "attachment" in download.headers["content-disposition"]

    # delete (not used yet → allowed)
    deleted = client.delete(f"/api/projects/{project_id}/attachments/{attachment_id}")
    assert deleted.status_code == 200
    remaining = client.get(f"/api/projects/{project_id}/attachments").json()
    assert all(item["attachment_id"] != attachment_id for item in remaining)


def test_api_download_unknown_attachment_returns_404(tmp_path: Path) -> None:
    client = _api_client(tmp_path)
    project_id = _create_project(client)
    response = client.get(f"/api/projects/{project_id}/attachments/does-not-exist/download")
    assert response.status_code == 404


def test_finalize_setup_selects_packs_from_attachment_text(tmp_path: Path) -> None:
    """Подбор доменных пакетов при отложенном setup учитывает ТЕКСТ вложений,
    а не только бизнес-запрос (разбор инцидента РТК: запрос «РТК», весь контекст
    — в приложенных файлах)."""
    client = _api_client(tmp_path)
    minimal_request = "Нужен пилот."
    rich = (
        "Нужен PoV по предиктивной аналитике оттока на ML. "
        "Источники: 1С и корпоративный портал. "
        "Нужны API-обновления, on-prem, персональные данные, BI и веб-интерфейс."
    )

    # Контроль: тот же скудный запрос БЕЗ вложений — подбор почти ничего не даёт.
    control = client.post(
        "/api/projects",
        json={
            "name": "control",
            "objective_ref": OBJECTIVE_REF,
            "request_text": minimal_request,
            "domain_pack_refs": [],
            "selection_provider": "stub",
        },
    )
    assert control.status_code == 200, control.text
    control_packs = set(control.json()["domain_pack_refs"])

    # 1. Создаём с отложенным setup и тем же скудным запросом.
    created = client.post(
        "/api/projects",
        json={
            "name": "deferred",
            "objective_ref": OBJECTIVE_REF,
            "request_text": minimal_request,
            "domain_pack_refs": [],
            "selection_provider": "stub",
            "defer_setup": True,
        },
    )
    assert created.status_code == 200, created.text
    body = created.json()
    assert body["setup_pending"] is True
    assert body["domain_pack_refs"] == []  # подбор отложен
    project_id = body["project_id"]

    # 2. Грузим вложение с доменными сигналами (синхронное извлечение).
    upload = client.post(
        f"/api/projects/{project_id}/attachments",
        files={"file": ("brief.txt", rich.encode("utf-8"), "text/plain")},
        data={"sync": "true"},
    )
    assert upload.status_code == 200, upload.text
    assert upload.json()["extraction_status"] == "succeeded"

    # 3. finalize → подбор по запросу + вложению.
    finalize = client.post(
        f"/api/projects/{project_id}/finalize-setup",
        json={"selection_provider": "stub"},
    )
    assert finalize.status_code == 200, finalize.text
    final_body = finalize.json()
    assert final_body["setup_pending"] is False
    final_packs = set(final_body["domain_pack_refs"])

    # Вложение реально повлияло на выбор: пакетов стало больше, чем по запросу.
    assert final_packs > control_packs
    assert "ml.predictive_analytics@1.0.0" in final_packs

    # Граф развёрнут только после finalize (есть задачи).
    graph = client.get(f"/api/projects/{project_id}/task-graph").json()
    assert graph["nodes"], "после finalize граф должен быть развёрнут"
