"""Сервис входных файлов-вложений проекта.

Загрузка, асинхронное извлечение текста (.pdf/.docx/текстовые), подача
извлечённого текста задачам через слой A (Position ``source="input"``, как
бизнес-запрос) и удаление-до-первого-использования.

Извлечение идёт в фоновом потоке (по образцу ``workflow_runner_service``):
загрузка отвечает быстро со статусом ``pending``, извлечение обновляет статус
позже. Graceful degradation: если текст извлечь нельзя — файл всё равно
сохранён и скачивается, помечается ``failed``/``unsupported`` и в контекст не
попадает.
"""

from __future__ import annotations

import threading
import uuid
from hashlib import sha256
from io import BytesIO
from pathlib import Path, PurePosixPath

from ..common.errors import ConflictError, PovGeneratorError
from ..common.serialization import utc_now_iso
from ..domain.attachments import AttachmentRecord, ExtractionStatus
from ..domain.positions import Position
from ..domain.project_knowledge import (
    RejectPositionPatch,
    UpsertPositionPatch,
)
from ..infrastructure.sqlite_runtime import SqliteRuntime

# Расширения, читаемые напрямую как UTF-8 текст.
_TEXT_EXTENSIONS = {".txt", ".md", ".json", ".csv"}
# Расширения с извлечением текста через библиотеки.
_DOC_EXTENSIONS = {".pdf", ".docx"}

# Разумные дефолты лимитов (§5.1.7 SPEC).
_MAX_FILE_BYTES = 25 * 1024 * 1024
_MAX_FILES_PER_PROJECT = 50

# Размер фрагмента извлечённого текста, попадающего в Position слоя A. Полный
# текст всегда лежит на диске; в положение кладём ограниченный фрагмент, чтобы
# не раздувать snapshot БД. При превышении — явный маркер усечения (не молча).
_POSITION_TEXT_CHAR_LIMIT = 20_000

# Префикс identifier'а положения, рождённого вложением. По нему context_service
# определяет, какие вложения реально вошли в контекст задачи.
ATTACHMENT_POSITION_PREFIX = "attachment."


def attachment_id_from_position_id(position_id: str) -> str | None:
    """Извлечь ``attachment_id`` из identifier'а связанного положения."""
    if not position_id.startswith(ATTACHMENT_POSITION_PREFIX):
        return None
    return position_id[len(ATTACHMENT_POSITION_PREFIX) :]


def _sanitize_filename(filename: str) -> str:
    """Базовая гигиена имени: без path traversal, без управляющих символов.

    Кириллица/юникод сохраняются. Хранение всё равно идёт под сгенерированным
    ``attachment_id`` — это имя используется только для отображения/скачивания.
    """
    # Отбросить любые компоненты пути (и posix, и windows-разделители).
    base = PurePosixPath(filename.replace("\\", "/")).name
    cleaned = "".join(ch for ch in base if ch.isprintable() and ch not in "\t\n\r").strip()
    cleaned = cleaned.strip(". ")
    if not cleaned:
        cleaned = "file"
    return cleaned[:255]


class AttachmentService:
    def __init__(
        self,
        runtime: SqliteRuntime,
        *,
        max_file_bytes: int = _MAX_FILE_BYTES,
        max_files_per_project: int = _MAX_FILES_PER_PROJECT,
    ) -> None:
        self._runtime = runtime
        self._max_file_bytes = max_file_bytes
        self._max_files_per_project = max_files_per_project

    @property
    def max_file_bytes(self) -> int:
        """Лимит размера одного файла (байт). Слой API использует его, чтобы
        ограничить чтение потока ДО материализации файла в память."""
        return self._max_file_bytes

    # --- загрузка -----------------------------------------------------------

    def upload(
        self,
        workspace: Path,
        project_id: str,
        *,
        filename: str,
        content: bytes,
        mime_type: str | None = None,
        extract_in_background: bool = True,
        purpose: str = "input",
    ) -> AttachmentRecord:
        """Сохранить файл и поставить извлечение текста (если формат поддержан).

        Лимиты enforced ДО сохранения; превышение → понятная ошибка.
        Повторная загрузка идентичного файла (то же содержимое и имя) —
        идемпотентна: возвращается уже существующая запись.
        """
        if not content:
            raise ConflictError("Пустой файл (0 байт) — сохранять нечего.")
        if len(content) > self._max_file_bytes:
            raise ConflictError(f"Файл превышает лимит размера ({len(content)} > {self._max_file_bytes} байт).")

        content_sha256 = sha256(content).hexdigest()
        safe_name = _sanitize_filename(filename)
        existing = self._runtime.list_attachments(workspace)
        # Дедупликация: точный повтор (то же содержимое И то же имя) не плодит
        # дубликат и не запускает повторное извлечение — возвращаем прежнюю
        # запись. Разное имя при том же содержимом считаем намеренно разными
        # файлами и не схлопываем.
        for prior in existing:
            if prior.sha256 == content_sha256 and prior.original_filename == safe_name:
                return prior
        if len(existing) >= self._max_files_per_project:
            raise ConflictError(f"Достигнут лимит количества вложений на проект ({self._max_files_per_project}).")

        extension = PurePosixPath(safe_name).suffix.lower()
        attachment_id = str(uuid.uuid4())
        storage_path = f"attachments/{attachment_id}{extension}"
        supported = extension in _TEXT_EXTENSIONS or extension in _DOC_EXTENSIONS
        status: ExtractionStatus = "pending" if supported else "unsupported"
        record = AttachmentRecord(
            attachment_id=attachment_id,
            project_id=project_id,
            original_filename=safe_name,
            mime_type=(mime_type or "application/octet-stream"),
            size_bytes=len(content),
            sha256=content_sha256,
            storage_path=storage_path,
            extraction_status=status,
            created_at=utc_now_iso(),
            extraction_error=(
                None if supported else f"Формат '{extension or '?'}' не поддержан для извлечения текста."
            ),
            purpose=purpose if purpose in {"input", "requisite"} else "input",
        )
        self._runtime.store_attachment(workspace, attachment=record, content=content)
        if supported:
            if extract_in_background:
                threading.Thread(
                    target=self._safe_extract,
                    args=(workspace, attachment_id),
                    daemon=True,
                ).start()
            else:
                # Синхронное извлечение: возвращаем обновлённую запись с
                # финальным статусом, а не pending-черновик до извлечения.
                return self.extract(workspace, attachment_id)
        return record

    # --- извлечение текста --------------------------------------------------

    def _safe_extract(self, workspace: Path, attachment_id: str) -> None:
        """Обёртка для фонового потока: извлечение не должно падать наружу."""
        try:
            self.extract(workspace, attachment_id)
        except Exception:  # noqa: BLE001 — фон, ошибки уже записаны в статус
            pass

    def extract(self, workspace: Path, attachment_id: str) -> AttachmentRecord:
        """Извлечь текст, записать его на диск и подать в слой A.

        Любая ошибка извлечения → статус ``failed``/``unsupported`` без падения
        (graceful degradation): файл остаётся скачиваемым, в контекст не идёт.
        """
        record = self._runtime.load_attachment(workspace, attachment_id)
        if record.is_deleted:
            return record
        try:
            content = self._runtime.load_attachment_content(workspace, attachment_id)
            text = self._extract_text(record, content)
        except Exception as error:  # noqa: BLE001
            return self._update(workspace, record, status="failed", error=str(error))

        if not text.strip():
            return self._update(
                workspace,
                record,
                status="failed",
                error="Текст не извлечён (пустой результат — возможно, скан без текстового слоя).",
            )

        # Извлечение (особенно крупного PDF) могло идти долго; за это время
        # вложение могли удалить (delete-до-использования). Перечитываем
        # запись и не «воскрешаем» удалённое: иначе оставили бы orphan-файл и
        # вернули бы Position в слой A.
        fresh = self._runtime.load_attachment(workspace, attachment_id)
        if fresh.is_deleted:
            return fresh

        text_ref = f"attachments/{attachment_id}.txt"
        (workspace / text_ref).write_text(text, encoding="utf-8")

        position_id = f"{ATTACHMENT_POSITION_PREFIX}{attachment_id}"
        self._runtime.apply_knowledge_patch(
            workspace,
            UpsertPositionPatch(self._build_position(position_id, fresh, text)),
            actor="attachment",
            reason=f"extracted text from attachment {fresh.original_filename}",
        )
        return self._update(
            workspace,
            fresh,
            status="succeeded",
            text_ref=text_ref,
            linked_position_id=position_id,
            error=None,
        )

    def _extract_text(self, record: AttachmentRecord, content: bytes) -> str:
        extension = PurePosixPath(record.storage_path).suffix.lower()
        if extension in _TEXT_EXTENSIONS:
            # NUL-байты — надёжный признак бинарного содержимого под текстовым
            # расширением. Не выдаём «текст» из мусора (errors="replace" дал бы
            # строку из replacement-символов), а помечаем файл как нечитаемый.
            if b"\x00" in content:
                raise PovGeneratorError(
                    "Файл выглядит бинарным (содержит NUL-байты) при текстовом расширении — "
                    "текст не извлечён."
                )
            return content.decode("utf-8", errors="replace")
        if extension == ".pdf":
            return self._extract_pdf(content)
        if extension == ".docx":
            return self._extract_docx(content)
        raise PovGeneratorError(f"Формат '{extension}' не поддержан для извлечения текста.")

    @staticmethod
    def _extract_pdf(content: bytes) -> str:
        from pypdf import PdfReader

        reader = PdfReader(BytesIO(content))
        return "\n".join((page.extract_text() or "") for page in reader.pages).strip()

    @staticmethod
    def _extract_docx(content: bytes) -> str:
        """Извлечь текст из .docx с сохранением таблиц и порядка блоков.

        Прежняя версия читала только ``document.paragraphs`` и теряла:
        таблицы (их ячейки — отдельная модель, не параграфы), колонтитулы,
        текст внутри гиперссылок и надписей (textbox). Здесь обходим тело
        документа В ПОРЯДКЕ следования (``iter_inner_content`` → параграфы и
        таблицы вперемешку), таблицы разворачиваем построчно (включая
        вложенные), плюс добираем колонтитулы. Каждый блок защищён
        try/except: один битый элемент не должен терять весь документ.
        """
        from docx import Document
        from docx.oxml.ns import qn
        from docx.table import Table
        from docx.text.paragraph import Paragraph

        document = Document(BytesIO(content))

        def para_text(paragraph: Paragraph) -> str:
            # Собираем все <w:t> (в т.ч. внутри гиперссылок и textbox, которые
            # paragraph.text пропускает). .iter рекурсивен по всем потомкам.
            parts = [node.text for node in paragraph._element.iter(qn("w:t")) if node.text]
            return "".join(parts).strip()

        def render_table(table: Table) -> list[str]:
            rows_out: list[str] = []
            for row in table.rows:
                cells_out: list[str] = []
                for cell in row.cells:
                    fragments: list[str] = []
                    for block in cell.iter_inner_content():
                        if isinstance(block, Paragraph):
                            text = para_text(block)
                            if text:
                                fragments.append(text)
                        elif isinstance(block, Table):
                            fragments.extend(render_table(block))  # вложенная таблица
                    cells_out.append(" ".join(fragments).replace("\n", " ").strip())
                if any(cells_out):  # пустые строки таблицы пропускаем
                    rows_out.append("| " + " | ".join(cells_out) + " |")
            return rows_out

        lines: list[str] = []
        for block in document.iter_inner_content():
            try:
                if isinstance(block, Paragraph):
                    text = para_text(block)
                    if text:
                        lines.append(text)
                elif isinstance(block, Table):
                    lines.extend(render_table(block))
            except Exception:  # noqa: BLE001 — битый блок не теряет весь файл
                continue

        # Колонтитулы (часто несут реквизиты/версии/подписи). Дедуп: один и тот
        # же колонтитул повторяется на каждой секции.
        seen: set[str] = set()
        for section in document.sections:
            for header_footer in (section.header, section.footer):
                try:
                    extra: list[str] = [para_text(p) for p in header_footer.paragraphs]
                    for table in header_footer.tables:
                        extra.extend(render_table(table))
                except Exception:  # noqa: BLE001
                    continue
                for item in extra:
                    if item and item not in seen:
                        seen.add(item)
                        lines.append(item)

        return "\n".join(lines).strip()

    def _build_position(self, position_id: str, record: AttachmentRecord, text: str) -> Position:
        statement_text = text
        if len(statement_text) > _POSITION_TEXT_CHAR_LIMIT:
            statement_text = (
                statement_text[:_POSITION_TEXT_CHAR_LIMIT].rstrip()
                + "\n\n… [текст файла обрезан для контекста; полный текст — в скачиваемом файле]"
            )
        statement = f"Входной файл «{record.original_filename}»:\n\n{statement_text}"
        return Position(
            identifier=position_id,
            type="fact",
            statement=statement,
            visibility="architectural",
            scope="global",
            source="input",
            taken_by="attachment",
            taken_at=utc_now_iso(),
            tags=("attachment", "input"),
        )

    # --- удаление -----------------------------------------------------------

    def delete(self, workspace: Path, attachment_id: str) -> None:
        """Удалить вложение — только пока оно не использовано в контексте.

        После удаления связанная Position слоя A снимается (reject): до
        использования она ещё ни на что не влияла.
        """
        record = self._runtime.load_attachment(workspace, attachment_id)
        if record.is_deleted:
            return
        if record.used_in_context:
            raise ConflictError(
                "Вложение уже использовано в контексте задачи — удаление запрещено "
                "ради воспроизводимости созданных артефактов."
            )
        if record.linked_position_id:
            self._runtime.apply_knowledge_patch(
                workspace,
                RejectPositionPatch(
                    position_id=record.linked_position_id,
                    reason=f"attachment {record.original_filename} deleted before use",
                ),
                actor="attachment",
                reason="attachment deleted before first use",
            )
        for relative in (record.storage_path, record.extracted_text_ref):
            if relative:
                (workspace / relative).unlink(missing_ok=True)
        self._update(
            workspace,
            record,
            status=record.extraction_status,
            error=record.extraction_error,
            is_deleted=True,
        )

    # --- helpers ------------------------------------------------------------

    def _update(
        self,
        workspace: Path,
        record: AttachmentRecord,
        *,
        status: ExtractionStatus,
        text_ref: str | None = None,
        linked_position_id: str | None = None,
        error: str | None = None,
        is_deleted: bool | None = None,
    ) -> AttachmentRecord:
        from dataclasses import replace

        updated = replace(
            record,
            extraction_status=status,
            extracted_text_ref=text_ref if text_ref is not None else record.extracted_text_ref,
            linked_position_id=(linked_position_id if linked_position_id is not None else record.linked_position_id),
            extraction_error=error,
            is_deleted=record.is_deleted if is_deleted is None else is_deleted,
        )
        return self._runtime.update_attachment(workspace, updated)
